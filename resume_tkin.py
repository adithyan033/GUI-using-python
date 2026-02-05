from tkinter import *
from tkinter import ttk, messagebox
from docx import Document
    
root = Tk()
root.title("Resume generator")

name_var = StringVar()
age_var = StringVar()
gender_var = StringVar()
qual_var = StringVar()
dob_var = StringVar()

def save_resume():
    if name_var.get() == "" or age_var == "" or gender_var == "" or qual_var == "" or dob_var == "":
        messagebox.showerror("Error", "All fields are required")
        return
    
    doc = Document()
    doc.add_heading("RESUME", level = 1)
    
    doc.add_paragraph(f"Name : {name_var.get()}")
    doc.add_paragraph(f"Age : {age_var.get()}")
    doc.add_paragraph(f"Gender : {gender_var.get()}")
    doc.add_paragraph(f"Qualification : {qual_var.get()}")
    doc.add_paragraph(f"Date of Birth : {dob_var.get()}")
    
    doc.save("resume.docx")
    messagebox.showinfo("Success", "Resume saved as resume.docx")
    
Label(root, text = "Resume Generator", font = ("arial", 16, "bold")).grid(row = 0, column = 0)
    
Label(root, text = "Name").grid(row = 1, column = 0)
Entry(root, textvariable = name_var).grid(row = 1, column = 1)

Label(root, text = "Age").grid(row = 2, column = 0)
Entry(root, textvariable = age_var).grid(row = 2, column = 1)

Label(root, text = "Gender").grid(row = 3, column = 0)
Radiobutton(root, text = "Male", variable = gender_var, value = "male").grid(row = 3, column = 1)
Radiobutton(root, text = "Female", variable = gender_var, value = "female").grid(row = 4, column = 1)
Radiobutton(root, text = "Other", variable = gender_var, value = "other").grid(row = 5, column = 1)

Label(root, text = "Date of birth (DD-MM-YYYY)").grid(row = 6, column = 0)
Entry(root, textvariable = dob_var).grid(row = 6, column = 1)

Label(root, text = "Qualification").grid(row = 7, column = 0)
Qualification_box = ttk.Combobox(root, textvariable = qual_var, values = ["MCA", "M.Sc Computer Science", "MBA"], state = "readonly").grid(row = 7, column = 1)

Button(root, text = "Save Resume", command = save_resume).grid(row = 9, column = 1)

root.mainloop()